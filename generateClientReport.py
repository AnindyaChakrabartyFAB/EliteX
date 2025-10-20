#!/usr/bin/env python3
"""
generateClientReport.py

Reads the latest elite_analysis_*.txt from logs/ (or a user-specified file),
parses key metrics, generates visualizations, and outputs a professional
blue-themed MS Word report.
"""

import os
import re
import json
import argparse
from pathlib import Path
from datetime import datetime

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION


LOGS_DIR = Path("logs")
OUTPUT_DIR = Path("Output")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def find_latest_log_file() -> Path | None:
    if not LOGS_DIR.exists():
        return None
    candidates = sorted(LOGS_DIR.glob("elite_analysis_*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0] if candidates else None


def extract_json_like_sections(text: str) -> list[dict]:
    """Extract JSON-like blocks from the log to recover tool outputs when present.
    This attempts to find braces-enclosed sections and parse them as JSON.
    """
    results = []
    brace_blocks = re.findall(r"\{[\s\S]*?\}", text)
    for block in brace_blocks:
        try:
            obj = json.loads(block)
            if isinstance(obj, dict):
                results.append(obj)
        except Exception:
            continue
    return results


def parse_key_metrics(text: str) -> dict:
    """Parse known sections and numeric cues from the consolidated log.
    Returns a dict with keys used by chart builders.
    """
    data: dict[str, any] = {
        "client_id": None,
        "sections": {},
        "metrics": {},
        "json_objects": [],
    }

    # Capture client id from header lines like: Client: 36XALFR
    m = re.search(r"Client:\s*([A-Za-z0-9_-]+)", text)
    if m:
        data["client_id"] = m.group(1)

    # Split sections by markers used in EliteX.py
    sections = re.split(r"\n={3,}.*?\n|\n-+\n", text)
    # Fallback: keep entire text as one section if small split
    if len(sections) < 2:
        sections = [text]

    # Store sections with a heuristic title
    for chunk in sections:
        title_match = re.search(r"^(?:===\s*)?([A-Z][A-Za-z ]+)(?:\s+for Client .*?)?\s*$", chunk, re.MULTILINE)
        title = title_match.group(1).strip() if title_match else f"Section {len(data['sections'])+1}"
        data["sections"][title] = chunk

    # Try to extract some core numbers with regex patterns
    patterns = {
        "total_holdings_value": r"total holdings? value[^0-9]*([0-9,.]+)",
        "holdings_count": r"holdings? count[^0-9]*([0-9,.]+)",
      	"total_credit_volume": r"total credit volume[^0-9]*([0-9,.]+)",
        "transaction_count": r"transaction count[^0-9]*([0-9,.]+)",
        "accounts_count": r"accounts? count[^0-9]*([0-9,.]+)",
        "six_month_average": r"six[- ]month average[^0-9]*([0-9,.]+)",
        "current_month_deposit": r"current month deposit[^0-9]*([0-9,.]+)",
        "comparison_ratio": r"comparison ratio[^0-9]*([0-9,.]+)",
        "total_alerts": r"total alerts[^0-9]*([0-9,.]+)",
    }
    lowered = text.lower()
    for key, pat in patterns.items():
        mm = re.search(pat, lowered)
        if mm:
            try:
                data["metrics"][key] = float(mm.group(1).replace(",", ""))
            except Exception:
                data["metrics"][key] = mm.group(1)

    # Extract JSON-like tool outputs to mine structured lists like transactions or holdings
    data["json_objects"] = extract_json_like_sections(text)
    return data


def build_dataframes(parsed: dict) -> dict:
    """From parsed JSON objects, assemble helpful DataFrames if available."""
    dfs: dict[str, pd.DataFrame] = {}
    # Holdings
    holdings_rows = []
    tx_rows = []
    alerts_rows = []

    for obj in parsed.get("json_objects", []):
        if isinstance(obj, dict):
            # Holdings
            if "current_holdings" in obj and isinstance(obj["current_holdings"], list):
                for h in obj["current_holdings"]:
                    if isinstance(h, dict):
                        holdings_rows.append(h)
            # Transactions
            for key in ["banking_transactions", "credit_transactions"]:
                if key in obj and isinstance(obj[key], list):
                    for t in obj[key]:
                        if isinstance(t, dict):
                            tx_rows.append(t)
            # AEDB Alerts
            # V7: renamed to risk_alerts
            alerts_list = obj.get("risk_alerts")
            if isinstance(alerts_list, list):
                for a in alerts_list:
                    if isinstance(a, dict):
                        alerts_rows.append(a)

    if holdings_rows:
        dfs["holdings"] = pd.DataFrame(holdings_rows)
    if tx_rows:
        dfs["transactions"] = pd.DataFrame(tx_rows)
    if alerts_rows:
        dfs["alerts"] = pd.DataFrame(alerts_rows)

    return dfs


def save_chart(fig, path: Path):
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches='tight')
    plt.close(fig)


def create_charts(parsed: dict, dfs: dict, out_dir: Path) -> dict:
    """Create blue-themed charts; return dict of image paths by name."""
    sns.set_theme(style="whitegrid")
    # Blue and red professional theme
    sns.set_palette(["#0B5ED7", "#DC3545", "#1D7BFF", "#E35D6A", "#4DA3FF"])  # Blue + Red

    images: dict[str, Path] = {}
    out_dir.mkdir(parents=True, exist_ok=True)

    # Pie: Transaction type distribution if available
    tx_df = dfs.get("transactions")
    if tx_df is not None and "transaction_type" in tx_df.columns:
        counts = tx_df["transaction_type"].astype(str).value_counts().head(8)
        if not counts.empty:
            fig, ax = plt.subplots(figsize=(6, 6))
            ax.pie(counts.values, labels=counts.index, autopct='%1.1f%%', startangle=140,
                   textprops={"color": "#0B1B3B", "fontsize": 10})
            ax.set_title("Transaction Type Distribution", color="#0B1B3B", fontsize=14)
            img_path = out_dir / "pie_transactions.png"
            save_chart(fig, img_path)
            images["pie_transactions"] = img_path

    # Bar: Top holdings by value if available
    holdings_df = dfs.get("holdings")
    if holdings_df is not None:
        value_col = None
        for candidate in ["investment_value_usd", "value", "amount"]:
            if candidate in holdings_df.columns:
                value_col = candidate
                break
        name_col = None
        for candidate in ["security_name", "name", "product_name"]:
            if candidate in holdings_df.columns:
                name_col = candidate
                break
        if value_col and name_col:
            tmp = holdings_df[[name_col, value_col]].copy()
            tmp[value_col] = pd.to_numeric(tmp[value_col], errors='coerce')
            tmp = tmp.dropna().sort_values(value_col, ascending=False).head(10)
            if not tmp.empty:
                fig, ax = plt.subplots(figsize=(7, 4))
                sns.barplot(x=value_col, y=name_col, data=tmp, ax=ax)
                ax.set_title("Top Holdings by Value", color="#0B1B3B", fontsize=14)
                ax.set_xlabel("Value (USD)", color="#0B1B3B")
                ax.set_ylabel("")
                for spine in ax.spines.values():
                    spine.set_color("#A3CEFF")
                img_path = out_dir / "bar_top_holdings.png"
                save_chart(fig, img_path)
                images["bar_top_holdings"] = img_path

    # Bar: AEDB alert levels if available
    alerts_df = dfs.get("alerts")
    if alerts_df is not None:
        if "risk_level" in alerts_df.columns:
            tmp = alerts_df.copy()
            tmp["risk_level"] = pd.to_numeric(tmp["risk_level"], errors='coerce')
            tmp = tmp.dropna(subset=["risk_level"]).sort_values("risk_level", ascending=False).head(10)
            if not tmp.empty:
                fig, ax = plt.subplots(figsize=(7, 4))
                sns.barplot(x="risk_level", y=("risk_name" if "risk_name" in tmp.columns else tmp.index), data=tmp, ax=ax)
                ax.set_title("Top AEDB Alerts by Risk Level", color="#0B1B3B", fontsize=14)
                ax.set_xlabel("Risk Level", color="#0B1B3B")
                ax.set_ylabel("")
                for spine in ax.spines.values():
                    spine.set_color("#A3CEFF")
                img_path = out_dir / "bar_alerts.png"
                save_chart(fig, img_path)
                images["bar_alerts"] = img_path

    return images


def add_heading(document: Document, text: str, level: int = 0):
    p = document.add_heading(text, level=level)
    for run in p.runs:
        # Level-based coloring: top headings blue, subheadings red
        if level <= 1:
            run.font.color.rgb = RGBColor(11, 94, 215)  # #0B5ED7 (blue)
        else:
            run.font.color.rgb = RGBColor(220, 53, 69)  # #DC3545 (red)
        run.font.bold = True
    return p


def add_paragraph(document: Document, text: str, size: int = 11):
    p = document.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
    return p


def build_word_report(parsed: dict, charts: dict, output_path: Path, logo_path: Path | None = None):
    doc = Document()

    # Page setup: landscape for better charts
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    try:
        # Professional margins
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    except Exception:
        pass

    # Compute usable width within margins (in inches)
    try:
        page_width_inches = section.page_width.inches  # type: ignore[attr-defined]
        left_margin_inches = section.left_margin.inches  # type: ignore[attr-defined]
        right_margin_inches = section.right_margin.inches  # type: ignore[attr-defined]
        usable_width_inches = max(1.0, page_width_inches - left_margin_inches - right_margin_inches)
    except Exception:
        usable_width_inches = 9.0  # sensible default for landscape

    def add_picture_fit(image_path: Path, desired_width_inches: float | None = None):
        if desired_width_inches is None:
            width_inches = min(usable_width_inches, 7.5)
        else:
            width_inches = min(usable_width_inches, desired_width_inches)
        try:
            pic = doc.add_picture(str(image_path), width=Inches(width_inches))
            # Center the image
            last_par = doc.paragraphs[-1]
            last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return pic
        except Exception:
            return None

    def add_caption(text: str):
        cap = doc.add_paragraph(text)
        cap_format = cap.runs[0].font if cap.runs else cap.add_run("").font
        cap_format.size = Pt(9)
        cap_format.color.rgb = RGBColor(90, 90, 90)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return cap

    # Set base font
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    except Exception:
        pass

    # Optional logo at the top-left
    if logo_path and logo_path.exists():
        add_picture_fit(logo_path, desired_width_inches=1.8)

    # Title
    title = f"Elite Client Strategy Report"
    subtitle = f"Client: {parsed.get('client_id') or 'Unknown'} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"

    tp = doc.add_paragraph()
    run = tp.add_run(title)
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 94, 215)
    run.bold = True
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sp = doc.add_paragraph()
    srun = sp.add_run(subtitle)
    srun.font.size = Pt(12)
    srun.font.color.rgb = RGBColor(67, 97, 238)  # lighter blue

    doc.add_paragraph("")

    # Narrative Executive Summary
    add_heading(doc, "Executive Summary", level=1)
    metrics = parsed.get("metrics", {})
    client_id = parsed.get("client_id") or "the client"
    summary_lines = []
    tv = metrics.get("total_holdings_value")
    tc = metrics.get("total_credit_volume")
    acc = metrics.get("accounts_count")
    dep_avg = metrics.get("six_month_average")
    curr_dep = metrics.get("current_month_deposit")
    alerts = metrics.get("total_alerts")

    summary_lines.append(
        f"This report presents a consolidated view of {client_id}, integrating portfolio, banking, and risk insights to guide clear, data-driven actions.")
    if tv:
        summary_lines.append(f"Current investment exposure totals approximately ${tv:,.0f}, with opportunities identified for optimization.")
    if tc:
        summary_lines.append(f"Credit utilization over recent periods amounts to ${tc:,.0f}, informing capacity and suitability for new facilities.")
    if acc:
        summary_lines.append(f"The client maintains {int(acc)} active accounts across the relationship, supporting diversified servicing needs.")
    if dep_avg and curr_dep is not None:
        try:
            ratio = metrics.get("comparison_ratio") or (float(curr_dep) / float(dep_avg) if float(dep_avg) > 0 else 0)
        except Exception:
            ratio = None
        if ratio is not None:
            summary_lines.append(
                f"Deposits this month are ${float(curr_dep):,.0f} versus a six‑month average of ${float(dep_avg):,.0f} (ratio {ratio:.2f}), informing short‑term liquidity posture.")
    if alerts:
        summary_lines.append(f"Risk review indicates {int(alerts)} AEDB alerts; the top items are highlighted in the risk section.")

    for line in summary_lines:
        add_paragraph(doc, line)

    # Charts
    add_heading(doc, "Visual Analytics", level=1)
    if charts:
        for label, path in charts.items():
            add_paragraph(doc, label.replace('_', ' ').title())
            add_picture_fit(path, desired_width_inches=7.2)
            add_caption(label.replace('_', ' ').title())
            doc.add_paragraph("")
    else:
        add_paragraph(doc, "No charts were generated due to limited structured data in the log.")

    # Raw analysis sections (compact)
    add_heading(doc, "Agent Analyses (Narrative Extract)", level=1)
    for title, content in parsed.get("sections", {}).items():
        add_heading(doc, title, level=2)
        # Transform extracted text into a narrative-style preview
        preview = content.strip()
        # Light normalization
        preview = re.sub(r"\n{3,}", "\n\n", preview)
        # Trim oversized chunks
        if len(preview) > 2400:
            preview = preview[:2400] + "\n...[additional details available in source log]"
        add_paragraph(doc, preview, size=10)
        doc.add_paragraph("")

    # Footer
    add_paragraph(doc, "\n— End of Report —", size=10)

    doc.save(str(output_path))


def find_logo_path(search_dir: Path) -> Path | None:
    """Try to locate a FAB logo image in the project folder."""
    candidates = []
    patterns = [
        "*fab*.*", "*first*abu*dhabi*.*", "*abu*dhabi*bank*.*", "*bank*logo*.*", "*logo*.*"
    ]
    exts = {".png", ".jpg", ".jpeg", ".bmp"}
    for pat in patterns:
        for p in search_dir.glob(pat):
            if p.suffix.lower() in exts:
                candidates.append(p)
    return candidates[0] if candidates else None


def main():
    parser = argparse.ArgumentParser(description="Generate MS Word report from Elite logs")
    parser.add_argument("--log", type=str, default="", help="Path to a specific elite_analysis_*.txt log file")
    parser.add_argument("--out", type=str, default="", help="Output .docx path; default in Output/")
    parser.add_argument("--logo", type=str, default="", help="Optional path to FAB logo image (png/jpg)")
    args = parser.parse_args()

    if args.log:
        log_path = Path(args.log)
        if not log_path.exists():
            raise FileNotFoundError(f"Log file not found: {log_path}")
    else:
        log_path = find_latest_log_file()
        if not log_path:
            raise FileNotFoundError("No elite_analysis_*.txt file found in logs/")

    with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    parsed = parse_key_metrics(text)
    dfs = build_dataframes(parsed)

    # Charts directory under Output/tmp_clientid_timestamp
    client_id = parsed.get("client_id") or "client"
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    charts_dir = OUTPUT_DIR / f"charts_{client_id}_{stamp}"
    charts = create_charts(parsed, dfs, charts_dir)

    # Output path
    if args.out:
        output_path = Path(args.out)
        if output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")
    else:
        output_path = OUTPUT_DIR / f"Elite_Client_Report_{client_id}_{stamp}.docx"

    # Resolve logo path
    logo_path = Path(args.logo).resolve() if args.logo else None
    if not logo_path or not logo_path.exists():
        # Try to locate a logo in project root
        logo_path = find_logo_path(Path.cwd())

    build_word_report(parsed, charts, output_path, logo_path)
    print(f"✅ Report generated: {output_path}")


if __name__ == "__main__":
    main()


